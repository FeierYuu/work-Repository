#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建带图片的俄语测试文档
用于验证文档转演示文稿功能中俄语文本和图片处理的正确性
"""

import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image


def create_russian_test_document_with_images(output_path):
    """
    创建包含图片的俄语测试docx文档
    
    Args:
        output_path: 输出文件路径
    """
    # 创建文档
    doc = Document()
    
    # 添加俄语标题
    doc.add_heading('Тестовый документ с изображениями', 0)
    
    # 添加俄语说明文本
    doc.add_paragraph('Это тестовый документ для проверки функциональности преобразования документов в презентации, содержащий следующее:')
    doc.add_paragraph('1. Заголовки и параграфы текста')
    doc.add_paragraph('2. Табличные данные')
    doc.add_paragraph('3. Математические формулы')
    doc.add_paragraph('4. Изображения')
    
    # 添加章节标题
    doc.add_heading('Первая часть: Пример текста', level=1)
    doc.add_paragraph('Это пример текста для проверки извлечения и преобразования текстового содержимого.')
    doc.add_paragraph('Текстовое содержимое, состоящее из нескольких абзацев, должно быть обработано правильно.')
    
    # 添加俄语数学公式说明和公式
    doc.add_heading('Примеры математических формул', level=1)
    doc.add_paragraph('Ниже приведена математическая формула:')
    doc.add_paragraph('$\\sum_{i=1}^{n} i = \\frac{n(n+1)}{2}$')
    
    # 添加俄语表格
    doc.add_heading('Пример таблицы данных', level=1)
    doc.add_paragraph('Вот простая таблица:')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя'
    hdr_cells[1].text = 'Возраст'
    hdr_cells[2].text = 'Город'
    
    # 添加表格行
    row_cells = table.add_row().cells
    row_cells[0].text = 'Иван'
    row_cells[1].text = '28'
    row_cells[2].text = 'Москва'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Анна'
    row_cells[1].text = '32'
    row_cells[2].text = 'Санкт-Петербург'
    
    # 添加图片标题（俄语）
    doc.add_heading('Примеры изображений', level=1)
    doc.add_paragraph('Ниже приведены тестовые изображения:')
    
    # 创建一个简单的测试图片
    try:
        # 创建一个简单的彩色图像
        image = Image.new('RGB', (400, 300), color=(73, 109, 137))
        
        # 保存到临时内存
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        # 添加图片到文档
        doc.add_picture(img_byte_arr, width=Inches(4.0))
        doc.add_paragraph('Это первое тестовое изображение - синее фоновое изображение')
        
    except Exception as e:
        print(f"Ошибка при создании изображения: {e}")
    
    # 添加第二张图片（不同大小）
    try:
        # 创建一个不同颜色的图像
        image2 = Image.new('RGB', (300, 400), color=(220, 20, 60))
        
        # 保存到临时内存
        img_byte_arr2 = io.BytesIO()
        image2.save(img_byte_arr2, format='PNG')
        img_byte_arr2.seek(0)
        
        # 添加图片到文档
        doc.add_paragraph('\nВторое тестовое изображение：')
        doc.add_picture(img_byte_arr2, width=Inches(3.0))
        doc.add_paragraph('Это второе тестовое изображение - красное фоновое изображение')
        
    except Exception as e:
        print(f"Ошибка при создании второго изображения: {e}")
    
    # 添加结尾文本（俄语）
    doc.add_heading('Конец документа', level=1)
    doc.add_paragraph('Это заключительная часть документа, которая будет содержаться на последней странице презентации.')
    
    # 保存文档
    doc.save(output_path)
    print(f"Тестовый документ создан: {output_path}")


def main():
    """
    主函数
    """
    # 设置输出路径
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_file = os.path.join(output_dir, 'russian_test_with_images.docx')
    
    # 创建测试文档
    create_russian_test_document_with_images(output_file)


if __name__ == "__main__":
    main()