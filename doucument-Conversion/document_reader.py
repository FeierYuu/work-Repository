#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档读取模块
支持读取doc、docx和pdf格式的文档
"""

import os
import zipfile
import shutil
import tempfile
from typing import Dict, List, Any
from docx import Document
import pdfplumber
import PyPDF2


class DocumentReader:
    """
    文档读取器类，支持多种格式的文档读取
    """
      
    def __init__(self):
          self.temp_dirs = []  # 跟踪所有创建的临时目录，以便后续清理
      
    def read(self, file_path: str) -> Dict[str, Any]:
        """
        读取文档内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext in ['.doc', '.docx']:
            return self._read_docx(file_path)
        elif file_ext == '.pdf':
            return self._read_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """
        读取docx文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 判断是否为标题（根据样式名称或字体大小）
                    is_heading = False
                    style_name = paragraph.style.name.lower()
                    if any(heading in style_name for heading in ['heading', '标题']):
                        is_heading = True
                    
                    # 检查字体大小（简单判断）
                    font_size = None
                    if paragraph.runs:
                        font_size = paragraph.runs[0].font.size
                    
                    content.append({
                        'text': text,
                        'type': 'heading' if is_heading else 'paragraph',
                        'font_size': font_size.pt if font_size else None
                    })
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 提取图片信息
            images_info = self.get_document_images_with_positions(file_path)
            
            return {
                'content': content,
                'tables': tables,
                'images': images_info,
                'format': 'docx',
                'file_path': file_path
            }
            
        except Exception as e:
            raise Exception(f"读取DOCX文件失败: {str(e)}")
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        读取PDF文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文档内容的字典
        """
        try:
            content = []
            tables = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        # 简单的段落分割
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            para = para.strip()
                            if para:
                                # 简单的标题判断（基于字体大小和位置）
                                # 这是一个简化的实现，实际应用可能需要更复杂的逻辑
                                lines = para.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if line:
                                        content.append({
                                            'text': line,
                                            'type': 'paragraph',  # 默认类型，后续分析会更新
                                            'page': page_num
                                        })
                    
                    # 提取表格
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        # 过滤空行
                        filtered_table = [row for row in table if any(cell for cell in row)]
                        if filtered_table:
                            tables.append(filtered_table)
            
            return {
                'content': content,
                'tables': tables,
                'format': 'pdf',
                'file_path': file_path,
                'page_count': len(pdf.pages) if 'pdf' in locals() else 0
            }
            
        except Exception as e:
            raise Exception(f"读取PDF文件失败: {str(e)}")
    
    def __init__(self):
        self.temp_dirs = []  # 跟踪所有创建的临时目录，以便后续清理
    
    def extract_images(self, file_path: str) -> List[str]:
        """
        从文档中提取图片
        
        Args:
            file_path: 文件路径
            
        Returns:
            List[str]: 临时图片文件路径列表
        """
        print(f"开始提取图片: {file_path}")
        temp_images = []
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f"文件格式: {file_ext}")
        
        if file_ext == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        images = page.images
                        print(f"在PDF第{page_num}页找到{len(images)}个图片")
                        for img_idx, img in enumerate(images):
                            # 这里可以保存图片到临时文件
                            # 简化实现，实际应用需要提取图片数据
                            print(f"  图片{img_idx+1}: {img}")
                            pass
            except Exception as e:
                print(f"提取PDF图片失败: {str(e)}")
                import traceback
                traceback.print_exc()
        
        elif file_ext in ['.docx']:
            try:
                print("处理DOCX文件，准备提取图片...")
                # 创建临时目录来存储提取的图片
                temp_dir = tempfile.mkdtemp()
                print(f"创建临时目录: {temp_dir}")
                self.temp_dirs.append(temp_dir)  # 保存临时目录路径，以便后续清理
                
                # 打开docx文件（实际上是zip文件）
                print(f"打开DOCX文件: {file_path}")
                with zipfile.ZipFile(file_path, 'r') as z:  
                    # 获取所有文件列表用于调试
                    all_files = z.namelist()
                    print(f"DOCX文件中包含{len(all_files)}个文件")
                    
                    # 获取所有图片文件
                    image_files = [name for name in all_files 
                                  if name.startswith('word/media/')]
                    print(f"找到{len(image_files)}个图片文件")
                    
                    # 提取每个图片文件
                    for img_idx, img_name in enumerate(image_files):
                        print(f"\n处理图片{img_idx+1}: {img_name}")
                        # 获取图片文件的扩展名
                        ext = os.path.splitext(img_name)[1]
                        print(f"图片扩展名: {ext}")
                        # 创建临时文件路径
                        temp_file_path = os.path.join(temp_dir, f'image_{img_idx}{ext}')
                        print(f"目标临时文件: {temp_file_path}")
                        
                        # 提取图片到临时文件
                        try:
                            print(f"提取图片数据到临时文件...")
                            with z.open(img_name) as source, open(temp_file_path, 'wb') as target:
                                shutil.copyfileobj(source, target)
                            
                            # 验证提取的文件
                            if os.path.exists(temp_file_path):
                                file_size = os.path.getsize(temp_file_path)
                                print(f"图片提取成功，文件大小: {file_size} 字节")
                                temp_images.append(temp_file_path)
                                print(f"已添加到临时图片列表: {temp_file_path}")
                            else:
                                print(f"警告: 图片文件创建失败: {temp_file_path}")
                        except Exception as img_err:
                            print(f"提取单个图片失败 {img_name}: {str(img_err)}")
                            import traceback
                            traceback.print_exc()
                
                print(f"\n图片提取完成，共提取{len(temp_images)}个图片")
            
            except Exception as e:
                print(f"提取DOCX图片失败: {str(e)}")
                import traceback
                traceback.print_exc()
                # 清理临时目录
                if 'temp_dir' in locals() and os.path.exists(temp_dir):
                    print(f"清理失败的临时目录: {temp_dir}")
                    shutil.rmtree(temp_dir)
                    if temp_dir in self.temp_dirs:
                        self.temp_dirs.remove(temp_dir)
        
        print(f"图片提取函数返回{len(temp_images)}个图片路径")
        return temp_images
    
    def cleanup_temp_files(self):
        """
        清理所有创建的临时目录
        """
        for temp_dir in self.temp_dirs:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                    print(f"已清理临时目录: {temp_dir}")
                except Exception as e:
                    print(f"清理临时目录失败 {temp_dir}: {str(e)}")
        self.temp_dirs = []
    
    def get_document_images_with_positions(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从文档中提取图片及其在文档中的位置信息
        
        Args:
            file_path: 文件路径
            
        Returns:
            List[Dict]: 包含图片路径和位置信息的列表
        """
        print(f"开始获取文档图片位置信息: {file_path}")
        images_info = []
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f"文件格式: {file_ext}")
        
        if file_ext in ['.docx']:
            try:
                print("开始提取图片及其位置信息...")
                # 提取所有图片
                print("调用extract_images提取所有图片...")
                temp_images = self.extract_images(file_path)
                print(f"extract_images返回{len(temp_images)}个图片")
                
                # 创建临时目录来解压docx文件
                temp_dir = tempfile.mkdtemp()
                print(f"创建临时解压目录: {temp_dir}")
                
                # 打开docx文件并解压
                print(f"解压DOCX文件到临时目录...")
                with zipfile.ZipFile(file_path, 'r') as z:
                    z.extractall(temp_dir)
                print("DOCX文件解压完成")
                
                # 解析document.xml来确定图片位置
                print("打开文档对象进行图片位置分析...")
                doc = Document(file_path)
                print(f"文档包含{len(doc.paragraphs)}个段落")
                current_paragraph = 0
                image_found_count = 0
                
                # 遍历所有段落，检测是否包含图片
                print("遍历段落查找图片位置...")
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    try:
                        # 检查段落是否包含图片
                        has_image = any(shape for shape in paragraph._element.xpath('.//a:blip') 
                                      for ns in paragraph._element.nsmap.values() if ns == 'http://schemas.openxmlformats.org/drawingml/2006/main')
                        
                        if has_image:
                            image_found_count += 1
                            print(f"在段落{para_idx}中发现图片")
                            # 这个段落包含图片
                            if temp_images and current_paragraph < len(temp_images):
                                image_path = temp_images[current_paragraph]
                                print(f"  关联图片: {image_path}")
                                images_info.append({
                                    'path': image_path,
                                    'paragraph_index': para_idx,
                                    'type': 'image',
                                    'position': para_idx  # 用于排序
                                })
                                current_paragraph += 1
                                print(f"  已添加到images_info，当前索引: {current_paragraph-1}")
                    except Exception as para_err:
                        print(f"分析段落{para_idx}时出错: {str(para_err)}")
                
                print(f"\n段落遍历完成，发现{image_found_count}个图片位置")
                print(f"已关联{len(images_info)}个图片，还有{len(temp_images) - current_paragraph}个未关联图片")
                
                # 如果还有未匹配的图片，添加到末尾
                if current_paragraph < len(temp_images):
                    print(f"处理{len(temp_images) - current_paragraph}个未关联的图片...")
                    while current_paragraph < len(temp_images):
                        image_path = temp_images[current_paragraph]
                        print(f"  添加未关联图片: {image_path}")
                        images_info.append({
                            'path': image_path,
                            'paragraph_index': len(doc.paragraphs),
                            'type': 'image',
                            'position': len(doc.paragraphs)
                        })
                        current_paragraph += 1
                
                print(f"\n图片位置信息收集完成，共{len(images_info)}个图片信息")
                # 打印images_info详情用于调试
                for idx, img_info in enumerate(images_info):
                    print(f"  图片{idx+1}: 路径={img_info['path']}, 段落索引={img_info['paragraph_index']}")
                
            except Exception as e:
                print(f"获取DOCX图片位置失败: {str(e)}")
                import traceback
                traceback.print_exc()
                # 清理临时目录
                if 'temp_dir' in locals() and os.path.exists(temp_dir):
                    print(f"清理临时解压目录: {temp_dir}")
                    shutil.rmtree(temp_dir)
        
        print(f"get_document_images_with_positions函数返回{len(images_info)}个图片信息")
        return images_info


if __name__ == "__main__":
    # 测试代码
    reader = DocumentReader()
    print("文档读取模块已创建")