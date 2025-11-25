#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档生成脚本 - 用于复现docx图片转pptx的问题
"""

import os
import sys
from docx import Document
from docx.shared import Inches

# 添加当前目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def create_test_docx_with_image(output_path):
    """创建一个包含图片的简单docx测试文档"""
    print(f"\n========== 创建测试docx文档 ==========")
    print(f"输出路径: {output_path}")
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        print(f"创建输出目录: {output_dir}")
        os.makedirs(output_dir)
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('测试文档', level=0)
    print("添加文档标题")
    
    # 添加一些文本
    doc.add_paragraph('这是一个测试文档，用于测试docx转pptx时图片显示问题。')
    doc.add_paragraph('文档包含一些俄语文本: Привет, мир!')
    print("添加测试文本")
    
    # 创建一个简单的测试图片（使用Base64编码的PNG图片）
    import io
    from PIL import Image, ImageDraw, ImageFont
    
    # 创建一个简单的红色测试图片
    print("创建测试图片...")
    img = Image.new('RGB', (400, 300), color='white')
    d = ImageDraw.Draw(img)
    
    # 尝试加载字体，如果失败则使用默认字体
    try:
        # 尝试使用系统字体
        font = ImageFont.truetype("Arial.ttf", 24) 
    except:
        font = ImageFont.load_default()
    
    # 在图片上绘制文本
    d.rectangle([(50, 50), (350, 250)], fill='red')
    d.text((100, 120), "测试图片", fill='white', font=font)
    d.text((100, 160), "Test Image", fill='white', font=font)
    d.text((100, 200), "Тестовое изображение", fill='white', font=font)
    
    # 保存图片到内存
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    print(f"测试图片创建成功: 400x300 像素")
    
    # 添加图片到docx
    print("将图片添加到docx文档...")
    doc.add_picture(img_byte_arr, width=Inches(4))
    print("图片添加成功")
    
    # 再添加一些文本
    doc.add_paragraph('图片下方的文本内容。')
    
    # 保存docx文件
    print(f"保存docx文件: {output_path}")
    doc.save(output_path)
    print(f"✓ docx文件保存成功")
    
    # 验证文件是否存在
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path)
        print(f"docx文件大小: {file_size} 字节")
        return True
    else:
        print(f"✗ 错误: docx文件未创建成功")
        return False

def main():
    """主函数"""
    print("\n===== 图片转换问题测试工具 =====")
    
    # 设置输出路径
    test_docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_image_document.docx")
    test_pptx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_image_presentation.pptx")
    
    # 1. 创建测试docx文档
    if not create_test_docx_with_image(test_docx_path):
        print("创建测试文档失败，退出程序")
        return 1
    
    print(f"\n\n========== 测试docx文档创建完成 ==========")
    print(f"文档路径: {test_docx_path}")
    print(f"\n请使用以下命令运行文档转演示文稿:")
    print(f"python document_to_presentation.py {test_docx_path} {test_pptx_path}")
    
    return 0

if __name__ == "__main__":
    sys.exit(main())